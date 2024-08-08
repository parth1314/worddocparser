from docx import Document
import re

def extract_text(document, heading):
    paragraphs = list(document.paragraphs)
    for i in range(len(paragraphs)):
        if paragraphs[i].style.name.startswith('Heading') and paragraphs[i].text == heading:
            text = []
            for j in range(i+1, len(paragraphs)):
                if not paragraphs[j].style.name.startswith('Heading'):
                    text.append(paragraphs[j].text)
                else:
                    break
            return ' '.join(text)
    return None


def insert_into_table_2_and_3(main_table_rows, sub_table_rows, cnx):
    cursor = cnx.cursor()

    
    for row in sub_table_rows:
        query = "INSERT INTO control_obj_row_1 (id, id_num, description) VALUES (%s, %s, %s)"
        
        match = re.match(r'(\w+)\.(\d+)(.*)', row)
        if match:
            xyz = match.group(1)
            a = match.group(2)
            description = match.group(3).strip()

            
            control_obj_base = re.match(r'([A-Za-z]+)', xyz).group(1)  
            control_obj_num = f"{control_obj_base}1.0"
            id_num = f"{xyz}.{a}"
            
            cursor.execute("SELECT control_obj_num FROM control_objectives WHERE control_obj_num = %s", (control_obj_num,))
            control_obj_record = cursor.fetchone()
            if control_obj_record:  
                values = (control_obj_num, id_num, description)
                cursor.execute(query, values)
                cnx.commit()
            else:
                print(f"Error: control_obj_num '{control_obj_num}' not found in control_objectives. Skipping insertion.")
    
    print("Records inserted successfully into control_obj_row_1")

    
    for row in main_table_rows:
        query = "INSERT INTO control_obj_data (id, criteria_num, description_of_controlcaseinc_controls, tests_by_service_auditor, results_of_controls_tests) VALUES (%s, %s, %s, %s, %s)"
    
        match = re.match(r'(\w+)(\d+)\.(\d+)\.(\d+)(.*)', row)
        if match:
            xyz = match.group(1)
            a = match.group(2)
            b = match.group(3)
            c = match.group(4)
            id_num = f"{xyz}{a}.{b}"
            criteria_num = f"{xyz}{a}.{b}.{c}"
            details = row.split('|', 3)  
        
        
        cursor.execute("SELECT id_num FROM control_obj_row_1 WHERE id_num = %s", (id_num,))
        if cursor.fetchone():  
            values = (id_num, *details)
            cursor.execute(query, values)
            cnx.commit()
        else:
            print(f"Error: id_num '{id_num}' not found in control_obj_row_1. Skipping insertion.")
    
    
    print("Records inserted successfully into control_obj_data")

    cnx.commit()
    cursor.close()
    cnx.close()


def docx_tables_to_txt(docx_path, txt_path):
    doc = Document(docx_path)
    with open(txt_path, 'w') as txt_file:
        for i,table in enumerate(doc.tables):
            if i == 0: continue 
            if i%2==0: txt_file.write("Table Start\n")
            for j, row in enumerate(table.rows):
                if j == 0 and i%2 != 0: continue 
                cell_texts = [cell.text.strip() + '|' for cell in row.cells]
                if j == 1 and i%2!=0:
                    txt_file.write("Sub-table Start\n")
                txt_file.write("\t".join(cell_texts) + "\n")
            if j == 1 and i%2 != 0: txt_file.write("Sub-Table End\n\n")
            else: txt_file.write("Table End\n\n")


def write_output(file_path, main_table_rows, sub_table_rows):
    with open(file_path, 'w') as file:
        file.write("Main Table Rows:\n")
        for row in main_table_rows:
            file.write(row + "\n")

        file.write("\nSub Table Rows:\n")
        for row in sub_table_rows:
            file.write(row + "\n")


def parse_txt_file_test(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()

    main_table_rows = []
    sub_table_rows = []
    is_in_sub_table = False
    is_in_main_table = False
    current_main_row = ""

    for line in lines:
        line = line.strip()

        if "Sub-table Start" in line:
            is_in_sub_table = True
            is_in_main_table = False
        elif "Sub-table End" in line:
            is_in_sub_table = False
        elif "Table Start" in line:
            is_in_main_table = True
            is_in_sub_table = False
        elif "Table End" in line:
            is_in_main_table = False
        else:
            if is_in_sub_table:
                
                parts = line.split('\t')
                
                sub_table_rows.append(parts[0].strip())
            elif is_in_main_table:
                if re.match(r'^\w+\.\d+\.\d+', line):  
                    if current_main_row:  
                        main_table_rows.append(current_main_row)
                    current_main_row = line
                else:
                    current_main_row += " " + line  

    return main_table_rows, sub_table_rows

def insert_intro_trust_services_criteria(doc_path, cnx):
    
    doc = Document(doc_path)
    introduction = extract_text(doc,'INTRODUCTION')
    trust_services_criteria = extract_text(doc,'TRUST SERVICES CRITERIA FOR SECURITY-RELATED CONTROLS, AND TESTS OF CONTROLS')
    cursor = cnx.cursor()
    doc = Document(doc_path)
    mySql_insert_query = """INSERT INTO intro_and_trust_services_criteria (introduction, trust_services_criteria) 
                            VALUES 
                            (%s, %s) """

    record = (introduction, trust_services_criteria)
    cursor.execute(mySql_insert_query, record)
    cnx.commit()
    print("Record inserted successfully into intro_and_trust_services_criteria table")

def extract_specific_headings(doc_path):
    doc = Document(doc_path)
    headings = []
    
    pattern = re.compile(r'^(\w+1\.0)\s+(.+)$')

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            match = pattern.match(paragraph.text)
            if match:
                part1 = match.group(1)
                part2 = match.group(2)
                headings.append((part1, part2))

    return headings


def insert_headings_into_db(headings, cnx):
    cursor = cnx.cursor()
    insert_query = "INSERT INTO control_objectives (control_obj_num, description) VALUES (%s, %s)"
    cursor.executemany(insert_query, headings)
    print("Records inserted successfully into control_objectives")
    cnx.commit()  
    cursor.close()